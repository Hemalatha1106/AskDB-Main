import os
import base64
import tempfile
import re
from datetime import datetime

# Optional imports with fallbacks
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentGenerator:
    """
    Generates reports in Markdown, PDF, and DOCX formats.
    """

    def __init__(self, title: str, creator: str = "AskDB", database_name: str = "Unspecified Database"):
        self.title = title
        self.creator = creator
        self.database_name = database_name
        self.generated_date = datetime.now().strftime("%B %d, %Y")

    def _clean_markdown_to_html(self, text: str) -> str:
        """
        Translates simple Markdown text to HTML tags supported by ReportLab's Paragraph.
        """
        if not text:
            return ""
        
        # Escape XML entities first
        html = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        # Convert bold (**text** or __text__)
        html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", html)
        html = re.sub(r"__(.*?)__", r"<b>\1</b>", html)
        
        # Convert italics (*text* or _text_)
        html = re.sub(r"\*(.*?)\*", r"<i>\1</i>", html)
        html = re.sub(r"_(.*?)_", r"<i>\1</i>", html)
        
        # Convert code inline (`code`)
        html = re.sub(r"`(.*?)`", r"<font face='Courier' color='#7c3aed'><b>\1</b></font>", html)
        
        # Convert newlines to breaks
        html = html.replace("\n", "<br/>")
        return html

    def generate_markdown(self, data: dict) -> str:
        """
        Generates standard Markdown representation of the report.
        """
        sections = data.get("sections", [])
        
        md = []
        md.append(f"# {self.title}")
        md.append(f"**Date:** {self.generated_date}  ")
        md.append(f"**Database:** {self.database_name}  ")
        md.append(f"**Prepared by:** {self.creator} (AskDB)  ")
        md.append("\n---\n")
        
        # Executive Summary
        md.append("## Executive Summary")
        md.append(data.get("executive_summary", "No executive summary provided."))
        md.append("\n---\n")
        
        # Report Sections
        md.append("## Report Sections")
        for idx, sec in enumerate(sections, 1):
            sec_title = sec.get("title", f"Section {idx}")
            md.append(f"### {idx}. {sec_title}")
            
            # User query
            md.append(f"**User Question:** *{sec.get('query', '')}*")
            md.append("")
            
            # Answer
            md.append("**Analysis Summary:**")
            md.append(sec.get("answer", ""))
            md.append("")
            
            # Optional SQL query
            if sec.get("sql_query") and sec.get("show_sql", True):
                md.append("**SQL Query Executed:**")
                md.append("```sql")
                md.append(sec.get("sql_query").strip())
                md.append("```")
                md.append("")
                
            # Chart/Plot
            if sec.get("plot"):
                md.append("**Visual Insight:**")
                md.append("*(Chart embedded in PDF/DOCX attachments)*")
                md.append(f"![{sec_title}](data:image/png;base64,{sec.get('plot')})")
                md.append("")
                
            # Key insights
            if sec.get("key_insights"):
                md.append("**Key Insights:**")
                md.append(sec.get("key_insights"))
                md.append("")
                
            md.append("\n---\n")
            
        # Overall Findings
        md.append("## Overall Findings")
        md.append(data.get("overall_findings", "No overall findings provided."))
        md.append("\n---\n")
        
        # Recommendations
        md.append("## Recommendations")
        md.append(data.get("recommendations", "No recommendations provided."))
        md.append("\n---\n")
        
        return "\n".join(md)

    def generate_pdf(self, data: dict, output_path: str):
        """
        Generates a styled, executive PDF report using ReportLab.
        """
        if not REPORTLAB_AVAILABLE:
            # Resilient fallback to basic text report if ReportLab is missing
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(self.generate_markdown(data))
            return

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=54, bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Custom elegant styles matching AskDB brand
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=26,
            leading=32,
            textColor=colors.HexColor('#0f1117'),
            alignment=0, # Left-aligned
            spaceAfter=20
        )
        
        meta_style = ParagraphStyle(
            'DocMeta',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#6b7280'),
            spaceAfter=30
        )
        
        h1_style = ParagraphStyle(
            'DocH1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#3b82f6'),
            spaceBefore=18,
            spaceAfter=12,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#111c3d'),
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10.5,
            leading=16,
            textColor=colors.HexColor('#0f1117'),
            spaceAfter=10
        )

        query_style = ParagraphStyle(
            'DocQuery',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=10,
            leading=15,
            textColor=colors.HexColor('#111c3d'),
            backColor=colors.HexColor('#f3f4f6'),
            borderColor=colors.HexColor('#e5e7eb'),
            borderWidth=1,
            borderPadding=10,
            spaceAfter=12,
            borderRadius=6
        )

        sql_style = ParagraphStyle(
            'DocSQL',
            parent=styles['Code'],
            fontName='Courier',
            fontSize=9,
            leading=13,
            textColor=colors.HexColor('#7c3aed'),
            backColor=colors.HexColor('#f9fafb'),
            borderColor=colors.HexColor('#e5e7eb'),
            borderWidth=1,
            borderPadding=8,
            spaceAfter=12,
            borderRadius=4
        )

        story = []

        # Cover/Header Metadata
        story.append(Paragraph(self.title, title_style))
        meta_html = f"<b>Generated Date:</b> {self.generated_date} | <b>Database:</b> {self.database_name}<br/><b>Prepared using:</b> {self.creator} (AskDB)"
        story.append(Paragraph(meta_html, meta_style))
        story.append(Spacer(1, 10))

        # Executive Summary
        story.append(Paragraph("Executive Summary", h1_style))
        exec_sum_html = self._clean_markdown_to_html(data.get("executive_summary", ""))
        story.append(Paragraph(exec_sum_html, body_style))
        story.append(Spacer(1, 15))

        # Sections
        story.append(Paragraph("Report Sections", h1_style))
        
        sections = data.get("sections", [])
        temp_files = [] # Track temp image files to clean up later

        for idx, sec in enumerate(sections, 1):
            sec_title = sec.get("title", f"Section {idx}")
            story.append(Paragraph(f"{idx}. {sec_title}", h2_style))
            
            # Question
            query_html = f"<b>Question:</b> {sec.get('query', '')}"
            story.append(Paragraph(query_html, query_style))
            
            # Answer/Summary
            ans_html = self._clean_markdown_to_html(sec.get("answer", ""))
            story.append(Paragraph(ans_html, body_style))
            
            # Optional SQL query
            if sec.get("sql_query") and sec.get("show_sql", True):
                sql_html = f"<b>SQL QUERY:</b><br/>{sec.get('sql_query').strip()}"
                # Replace newlines with breaks in SQL
                sql_html = sql_html.replace("\n", "<br/>").replace(" ", "&nbsp;")
                story.append(Paragraph(sql_html, sql_style))
                
            # Base64 Chart image
            if sec.get("plot"):
                try:
                    img_data = base64.b64decode(sec.get("plot"))
                    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    temp_img.write(img_data)
                    temp_img.close()
                    temp_files.append(temp_img.name)
                    
                    # Create and scale image flowable
                    pdf_img = Image(temp_img.name, width=400, height=220)
                    pdf_img.hAlign = 'CENTER'
                    story.append(Spacer(1, 5))
                    story.append(pdf_img)
                    story.append(Spacer(1, 5))
                except Exception as img_err:
                    story.append(Paragraph(f"<i>(Failed to render visualization chart: {img_err})</i>", body_style))
            
            # Insights
            if sec.get("key_insights"):
                insights_html = f"<b>Key Insights:</b><br/>{self._clean_markdown_to_html(sec.get('key_insights'))}"
                story.append(Paragraph(insights_html, body_style))
                
            story.append(Spacer(1, 15))

        # Overall Findings
        story.append(Paragraph("Overall Findings", h1_style))
        findings_html = self._clean_markdown_to_html(data.get("overall_findings", ""))
        story.append(Paragraph(findings_html, body_style))
        story.append(Spacer(1, 15))

        # Recommendations
        story.append(Paragraph("Recommendations", h1_style))
        recs_html = self._clean_markdown_to_html(data.get("recommendations", ""))
        story.append(Paragraph(recs_html, body_style))

        # Build PDF
        try:
            doc.build(story)
        finally:
            # Clean up temp image files
            for file_path in temp_files:
                try:
                    if os.path.exists(file_path):
                        os.unlink(file_path)
                except:
                    pass

    def generate_docx(self, data: dict, output_path: str):
        """
        Generates a styled Microsoft Word Document using python-docx.
        """
        if not DOCX_AVAILABLE:
            # Resilient fallback to Markdown as text file
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(self.generate_markdown(data))
            return

        doc = docx.Document()

        # Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Style colors
        color_primary = RGBColor(59, 130, 246)   # Blue
        color_secondary = RGBColor(17, 28, 61)  # Dark Blue/Navy
        color_muted = RGBColor(107, 114, 128)   # Grey

        # Add Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(self.title)
        title_run.bold = True
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = color_secondary
        title_p.paragraph_format.space_after = Pt(4)

        # Add Metadata subtitle
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(f"Generated Date: {self.generated_date} | Database: {self.database_name}\nPrepared using: {self.creator} (AskDB)")
        meta_run.font.name = 'Arial'
        meta_run.font.size = Pt(9.5)
        meta_run.font.italic = True
        meta_run.font.color.rgb = color_muted
        meta_p.paragraph_format.space_after = Pt(24)

        def add_heading(text, level, space_before=18, space_after=8):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Arial'
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = color_primary
            else:
                run.font.size = Pt(13)
                run.font.color.rgb = color_secondary
            return p

        def add_body(text, space_after=8):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            return p

        def add_callout(text, prefix="Question: "):
            # We construct a table with 1 cell for a nice callout box
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            
            cell = table.cell(0, 0)
            
            # Apply cell background shading to light gray
            shading_elm = parse_xml(r'<w:shd {} w:fill="F3F4F6"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Left border thick gray accent
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="24" w:space="0" w:color="D1D5DB"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>'.format(nsdecls('w')))
            tcPr.append(tcBorders)

            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.right_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            r_prefix = p.add_run(prefix)
            r_prefix.bold = True
            r_prefix.font.name = 'Calibri'
            r_prefix.font.size = Pt(10.5)
            r_prefix.font.color.rgb = color_secondary
            
            r_text = p.add_run(text)
            r_text.italic = True
            r_text.font.name = 'Calibri'
            r_text.font.size = Pt(10.5)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6) # spacer below table

        def add_sql_block(sql_query):
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Inches(6.5)
            cell = table.cell(0, 0)
            
            # Background shading
            shd = parse_xml(r'<w:shd {} w:fill="F9FAFB"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shd)
            
            # Thin borders all around
            tcBorders = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/><w:top w:val="single" w:sz="4" w:color="E5E7EB"/><w:right w:val="single" w:sz="4" w:color="E5E7EB"/><w:bottom w:val="single" w:sz="4" w:color="E5E7EB"/></w:tcBorders>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(tcBorders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.left_indent = Inches(0.1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(sql_query.strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(124, 58, 237) # Purple
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Executive Summary
        add_heading("Executive Summary", level=1)
        add_body(data.get("executive_summary", ""))

        # Sections
        add_heading("Report Sections", level=1)
        
        report_sections = data.get("sections", [])
        temp_files = []

        for idx, sec in enumerate(report_sections, 1):
            sec_title = sec.get("title", f"Section {idx}")
            add_heading(f"{idx}. {sec_title}", level=2)
            
            # Callout for question
            add_callout(sec.get("query", ""), prefix="Question: ")
            
            # Body Answer
            add_body(sec.get("answer", ""))
            
            # SQL block
            if sec.get("sql_query") and sec.get("show_sql", True):
                add_body("SQL Query Executed:", space_after=2)
                add_sql_block(sec.get("sql_query"))
                
            # Chart
            if sec.get("plot"):
                try:
                    img_data = base64.b64decode(sec.get("plot"))
                    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    temp_img.write(img_data)
                    temp_img.close()
                    temp_files.append(temp_img.name)
                    
                    # Insert in Word
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.space_before = Pt(6)
                    p.add_run().add_picture(temp_img.name, width=Inches(5.0))
                except Exception as docx_img_err:
                    add_body(f"(Failed to insert chart image: {docx_img_err})")
                    
            # Insights
            if sec.get("key_insights"):
                add_body(f"Key Insights:", space_after=2)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(sec.get("key_insights"))
                run.italic = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

        # Overall Findings
        add_heading("Overall Findings", level=1)
        add_body(data.get("overall_findings", ""))

        # Recommendations
        add_heading("Recommendations", level=1)
        add_body(data.get("recommendations", ""))

        # Save Document
        try:
            doc.save(output_path)
        finally:
            # Clean up temp image files
            for file_path in temp_files:
                try:
                    if os.path.exists(file_path):
                        os.unlink(file_path)
                except:
                    pass
